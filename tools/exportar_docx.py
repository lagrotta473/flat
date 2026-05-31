"""
Converte um arquivo Markdown da knowledge_base para .docx (Word).

Uso:
    python tools/exportar_docx.py <arquivo.md> [--saida <destino.docx>]

Exemplos:
    python tools/exportar_docx.py knowledge_base/clausula_8_operacao/_overview.md
    python tools/exportar_docx.py knowledge_base/clausula_8_operacao/_overview.md --saida relatorio_operacao.docx

Saída padrão:
    Mesmo diretório do arquivo de entrada, com extensão .docx

Dependências (instalar via pip):
    pip install python-docx
"""

import argparse
import re
import sys
from pathlib import Path


def _md_para_docx(caminho_md: Path, caminho_saida: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        print(
            "Erro: python-docx não instalado.\nExecute: pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)

    doc = Document()

    # Estilo base
    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    linhas = caminho_md.read_text(encoding="utf-8").splitlines()
    em_tabela = False
    cabecalho_tabela: list[str] = []
    linhas_tabela: list[list[str]] = []
    em_lista = False

    def _flush_tabela():
        nonlocal em_tabela, cabecalho_tabela, linhas_tabela
        if not cabecalho_tabela:
            return
        n_cols = len(cabecalho_tabela)
        tabela = doc.add_table(rows=1 + len(linhas_tabela), cols=n_cols)
        tabela.style = "Table Grid"
        # Cabeçalho
        for i, cel in enumerate(cabecalho_tabela):
            cell = tabela.cell(0, i)
            cell.text = cel.strip()
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cel.strip())
            run.bold = True
        # Linhas
        for r, linha in enumerate(linhas_tabela):
            for c, cel in enumerate(linha):
                if c < n_cols:
                    tabela.cell(r + 1, c).text = cel.strip()
        doc.add_paragraph()
        em_tabela = False
        cabecalho_tabela = []
        linhas_tabela = []

    for linha in linhas:
        # Tabela Markdown
        if linha.startswith("|"):
            celulas = [c for c in linha.split("|")[1:-1]]
            # Linha separadora (---|---|)
            if all(re.match(r"[-: ]+$", c) for c in celulas):
                continue
            if not em_tabela:
                em_tabela = True
                cabecalho_tabela = celulas
            else:
                linhas_tabela.append(celulas)
            continue
        elif em_tabela:
            _flush_tabela()

        # Títulos
        if linha.startswith("# "):
            doc.add_heading(linha[2:].strip(), level=1)
        elif linha.startswith("## "):
            doc.add_heading(linha[3:].strip(), level=2)
        elif linha.startswith("### "):
            doc.add_heading(linha[4:].strip(), level=3)
        elif linha.startswith("#### "):
            doc.add_heading(linha[5:].strip(), level=4)

        # Listas
        elif re.match(r"^[-*+] ", linha):
            p = doc.add_paragraph(style="List Bullet")
            _adicionar_runs(p, linha[2:].strip())
        elif re.match(r"^\d+\. ", linha):
            p = doc.add_paragraph(style="List Number")
            _adicionar_runs(p, re.sub(r"^\d+\. ", "", linha).strip())

        # Linha em branco
        elif linha.strip() == "":
            doc.add_paragraph()

        # Linha horizontal
        elif re.match(r"^---+$", linha.strip()):
            doc.add_paragraph("─" * 40)

        # Parágrafo normal
        else:
            p = doc.add_paragraph()
            _adicionar_runs(p, linha.strip())

    if em_tabela:
        _flush_tabela()

    doc.save(str(caminho_saida))


def _adicionar_runs(paragrafo, texto: str) -> None:
    """Processa negrito (**texto**) e itálico (*texto*) no parágrafo."""
    # Divide por marcadores de formatação preservando os delimitadores
    partes = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            run = paragrafo.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith("*") and parte.endswith("*"):
            run = paragrafo.add_run(parte[1:-1])
            run.italic = True
        elif parte.startswith("`") and parte.endswith("`"):
            run = paragrafo.add_run(parte[1:-1])
            run.font.name = "Courier New"
        elif parte:
            paragrafo.add_run(parte)


def main():
    parser = argparse.ArgumentParser(
        description="Exporta arquivo Markdown para .docx (Word)."
    )
    parser.add_argument("markdown", type=Path, help="Arquivo .md a converter")
    parser.add_argument(
        "--saida",
        type=Path,
        default=None,
        help="Caminho de saída do .docx. Padrão: mesmo diretório do .md",
    )
    args = parser.parse_args()

    if not args.markdown.exists():
        print(f"Erro: arquivo não encontrado: {args.markdown}", file=sys.stderr)
        sys.exit(1)

    saida = args.saida or args.markdown.with_suffix(".docx")
    _md_para_docx(args.markdown, saida)
    print(f"✓ Exportado: {saida}")


if __name__ == "__main__":
    main()
